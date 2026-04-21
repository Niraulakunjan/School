from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_manual():
    doc = Document()

    # Title Page
    title = doc.add_heading('SajiloSchool User Manual', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Guide for School Principals and Staff')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    version = doc.add_paragraph('Version 1.0 | April 2026')
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "SajiloSchool is a modern, multi-tenant School Management System designed to streamline "
        "academic, administrative, and financial operations. This manual provides a step-by-step guide "
        "to using the platform effectively."
    )

    # 2. Getting Started
    doc.add_heading('2. Getting Started', level=1)
    doc.add_heading('2.1 Logging In', level=2)
    doc.add_paragraph(
        "Access your school's unique URL (e.g., yourschool.sajiloschool.com). "
        "Enter your username and password provided by the administrator. The system "
        "uses secure JWT authentication to ensure your data stays private."
    )
    
    doc.add_heading('2.2 The Dashboard', level=2)
    doc.add_paragraph(
        "Upon logging in, you will see the Dashboard Analytics. Here you can find "
        "quick stats on student enrollment, total collections, and upcoming exams."
    )

    # 3. Academic Management
    doc.add_heading('3. Academic Management', level=1)
    doc.add_paragraph(
        "Setting up your academic structure is the first step in using SajiloSchool."
    )
    
    doc.add_heading('3.1 Classes and Sections', level=2)
    doc.add_paragraph(
        "Define your grades (e.g., 1, 2, 11, 12) and their respective sections (A, B, C). "
        "Each class can be associated with a specific faculty if applicable."
    )
    
    doc.add_heading('3.2 Subjects and Assignments', level=2)
    doc.add_paragraph(
        "Create subjects and assign them to classes using the Class-Subject Mapping tool. "
        "You can define Full Marks, Pass Marks, and Credit Hours for each assignment."
    )

    # 4. Student Management
    doc.add_heading('4. Student Management', level=1)
    
    doc.add_heading('4.1 Admissions Workflow', level=2)
    doc.add_paragraph(
        "Adding a new student is simple. Navigate to the Admission form and fill in:"
    )
    records = [
        ('Personal Info', 'Name, DOB, Gender, Blood Group'),
        ('Academic Info', 'Grade, Section, Symbols, Registration No.'),
        ('Guardian Info', 'Names, Contact Details, Relations'),
        ('Documents', 'Birth Certificates, Records, Photos')
    ]
    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Details to Capture'
    for cat, det in records:
        row_cells = table.add_row().cells
        row_cells[0].text = cat
        row_cells[1].text = det

    doc.add_heading('4.2 Attendance Tracking', level=2)
    doc.add_paragraph(
        "Teachers can mark daily attendance. The system supports various statuses: "
        "Present, Absent, Late, and Holiday."
    )

    # 5. Examination Management
    doc.add_heading('5. Examination Management', level=1)
    
    doc.add_heading('5.1 Exam Setup & Routines', level=2)
    doc.add_paragraph(
        "Create examination cycles (e.g., First Term, Final Term). For each exam, "
        "you can generate a professional schedule (Routine) that maps subjects to dates and times."
    )
    
    doc.add_heading('5.2 Mark Entry & Ledger', level=2)
    doc.add_paragraph(
        "Staff can enter marks directly into the system using the Mark Ledger. "
        "The system separates Theory and Practical marks for accurate calculation."
    )
    
    doc.add_heading('5.3 Grade Sheets (Certificates)', level=2)
    doc.add_paragraph(
        "Once marks are entered, the system automatically calculates GPAs and Grades. "
        "Professional Grade Sheets can be printed in bulk or individually for every student."
    )

    # 6. Finance & Fee Collection
    doc.add_heading('6. Finance & Fees', level=1)
    
    doc.add_heading('6.1 Fee Structures', level=2)
    doc.add_paragraph(
        "Define fee packages for each class. You can include items like Tuition Fees, "
        "Admission Fees, and Bus Fees."
    )
    
    doc.add_heading('6.2 Payment Collection', level=2)
    doc.add_paragraph(
        "Record payments easily. The system tracks balances, outstanding dues, "
        "and provides instant receipts for parents."
    )

    # 7. User Flows
    doc.add_heading('7. Common User Flows', level=1)
    
    doc.add_heading('7.1 New Student Cycle', level=2)
    doc.add_paragraph(
        "Admission -> Section Assignment -> Fee Enrollment -> Attendance Tracking"
    )

    doc.add_heading('7.2 Examination Cycle', level=2)
    doc.add_paragraph(
        "Create Exam -> Map Routines -> Mark Entry -> View Ledger -> Print Certificates"
    )

    # Footer
    section = doc.sections[0]
    footer = section.footer
    doc.add_paragraph('\n' * 2)
    doc.add_paragraph('Prepared for SajiloSchool Partners | 2026')

    # Save
    doc.save('SajiloSchool_User_Manual.docx')
    print("Manual generated: SajiloSchool_User_Manual.docx")

if __name__ == "__main__":
    generate_manual()
