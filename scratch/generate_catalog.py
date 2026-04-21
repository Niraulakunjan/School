from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_feature_section(doc, title, description, details):
    doc.add_heading(title, level=1)
    doc.add_paragraph(description)
    for head, text in details:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{head}: ")
        run.bold = True
        p.add_run(text)

def generate_catalog():
    doc = Document()

    # Styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Cover Page
    title = doc.add_heading('SajiloSchool — The Complete Feature Catalog', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 3)
    intro = doc.add_paragraph(
        "A detailed plain-English guide to every single helper tool inside your school system. "
        "No Computer Language. Just simple explanations for school owners and staff."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # 1. THE DIGITAL MASTER KEY (Multi-Tenancy)
    add_feature_section(doc, "1. Your Own Private Digital School", 
        "The system creates a private space just for your school. It’s like having your own "
        "locked office inside a big building.", 
        [
            ("Private Rooms", "Your school’s information is kept in its own digital 'room'. No other school can ever see your students or money records."),
            ("Branded Web Address", "You get a unique website address (like 'myschool.sajiloschool.com'). It feels professional and is easy for parents to remember."),
            ("Secret Locks", "Only people you allow can enter. Even the internal system names are hidden to keep hackers away.")
        ]
    )

    # 2. THE MASTER CONTROL ROOM (Settings)
    add_feature_section(doc, "2. Setting Up Your School Year", 
        "General settings that help the computer understand how your school operates.", 
        [
            ("The School Calendar", "Set your start and end dates for the year. This helps the computer know when to start fresh and when to finish."),
            ("Custom Logos", "Upload your school’s logo. It will automatically show up on every report card and receipt."),
            ("Role Management", "Decide who is a Teacher, a Student, or a Principal. Each person only sees the buttons they need.")
        ]
    )

    # 3. CLASSROOM BLUEPRINTS (Academics)
    add_feature_section(doc, "3. Grades, Sections & Subjects", 
        "Tell the computer what happens inside your classrooms.", 
        [
            ("Grades and Sections", "Mark down every grade from 1 to 12. You can even have Sections like 'A' and 'B' for busy schools."),
            ("Subject Master List", "Write down every subject you teach. You only have to do this once!"),
            ("Subject Mapping", "Tell the computer which subject belongs to which grade. For example, 'Grade 10 has Science' but 'Grade 1 has Drawing'."),
            ("Weightage (Credit Hours)", "Tell the computer how 'heavy' a subject is. It uses this to calculate the GPA (the final score) correctly.")
        ]
    )

    # 4. THE STUDENT'S DIGITAL FOLDER (Detailed SIS)
    add_feature_section(doc, "4. Recording Every Detail About Each Child", 
        "The system stores 10 times more detail than a normal paper file ever could.", 
        [
            ("Personal Statistics", "Full names, exact birthdays (English & Nepali), Gender, and Blood Group for emergencies."),
            ("Health & Physicals", "Record current Height and Weight to see how the child is growing over the year."),
            ("Home Information", "Permanent and current addresses so you know where every child lives."),
            ("Family Tree", "Names of Fathers, Mothers, and Guardians, along with their mobile numbers and jobs."),
            ("Digital File Cabinet", "The system has a place for you to upload photos of Birth Certificates, Passports, and previous class records."),
            ("Roll Numbers & IDs", "Every child gets an ID number that doesn't change, and a roll number for their current class.")
        ]
    )

    # 5. THE STAFF & TEACHER DIRECTORY
    add_feature_section(doc, "5. Managing Your Team", 
        "Keep track of the people who teach your students.", 
        [
            ("Teacher Files", "Record their qualifications (Degrees) and which department they work in."),
            ("Work Tracking", "Keep track of when they joined the school and their employee ID numbers.")
        ]
    )

    # 6. THE EXAM ENGINE & REPORT CARD FACTORY
    add_feature_section(doc, "6. From the First Day of Exams to the Final Report Card", 
        "The system handles the hardest part of school life automatically.", 
        [
            ("Exam Cycles", "Set up 'First Term' or 'Annual' exams. The system keeps them separate."),
            ("The Exam Calendar", "Click buttons to make a schedule. It tells students precisely when and where their tests are."),
            ("Theory vs Practical", "If a subject has a lab part (Science), you can put those marks separately. The computer adds them up for you."),
            ("The Big Mark Table", "One screen for teachers to type in marks. No more calculating on paper!"),
            ("The GPA Calculator", "The computer does all the math. It calculates the final GPA (e.g., 3.04) and 'Success' or 'Failed' status in a split second."),
            ("Beautiful Printing", "One-click printing for Grade Sheets. They look professional, are perfectly aligned, and have your logo on top.")
        ]
    )

    # 7. THE SCHOOL BANK (Finance)
    add_feature_section(doc, "7. Keeping Track of Every Penny", 
        "A clear view of income without any accounting errors.", 
        [
            ("Price Lists", "Create 'Fee Packages' for each class. You can have tuition fees, lab fees, and even bus fees."),
            ("Monthly Billing", "The computer automatically makes a 'Bill' for every child at the start of the month."),
            ("Scholarships & Discounts", "Give 50% off or a flat discount easily. The computer subtracts it from their bill automatically."),
            ("Digital Receipts", "When a parent pays, type the amount and the computer gives an instant receipt. It tracks exactly how much balance is left."),
            ("Expense Log", "Record money you spend on salaries or stationery so you know where your school's money is going.")
        ]
    )

    # 8. BUSES & TRANSPORT
    add_feature_section(doc, "8. Managing School Buses", 
        "Ensuring every child gets home safely.", 
        [
            ("Bus Route Setup", "Create routes and set the price for each one."),
            ("Link Students to Buses", "Assign students to their bus stops. The system automatically adds the bus fee to their monthly bill.")
        ]
    )

    # 9. ATTENDANCE & TRACKING
    add_feature_section(doc, "9. Daily Attendance Tracking", 
        "Knowing where every student is, every day.", 
        [
            ("Simple Markings", "Teachers can quickly mark Present or Absent on their phones or computers."),
            ("Attendance Records", "The system keeps a long memory. You can see a child's attendance for the whole year at once.")
        ]
    )

    # Footer
    doc.add_page_break()
    doc.add_heading("Why SajiloSchool is Better than Paper?", level=1)
    doc.add_paragraph("1. It never forgets a mark or a fee payment.")
    doc.add_paragraph("2. It does all the math correctly. No more calculator mistakes!")
    doc.add_paragraph("3. It saves hours of time. What used to take days (Report Cards) now takes seconds.")
    doc.add_paragraph("4. It keeps your school data safe. Paper files can burn or get lost; digital files are safe forever.")

    doc.save('SajiloSchool_Complete_Feature_Catalog.docx')
    print("Catalog generated: SajiloSchool_Complete_Feature_Catalog.docx")

if __name__ == "__main__":
    generate_catalog()
